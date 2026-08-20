#!/usr/bin/python
# -*- coding:utf-8 -*-
"""
MinerU-OCR WebUI (适配 llama.cpp 后端)
利用 llama-server (端口 50600) 进行极速推理，保留原版 PDF 分页、表格导出功能
"""
import gradio as gr
import os
import time
import tempfile
import gc
import shutil
import fitz
import re
import base64
import io
import requests
from PIL import Image
from docx import Document

# ==================== 配置（适配你的 llama.cpp 环境） ====================
LLAMA_API_URL = "http://127.0.0.1:50600/v1/chat/completions"  # 你日志里的端口
MODEL_NAME = "MinerU2.5-Pro"                                   # 你日志里的模型名
MAX_TOKENS = 2048
TEMPERATURE = 0.0

# ==================== 临时文件追踪 ====================
_temp_dirs = []

def cleanup_temp_dirs():
    for d in _temp_dirs:
        try:
            shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass
    _temp_dirs.clear()

# ==================== 辅助函数 ====================

def encode_image(image_path):
    """读取本地图片并转为 Base64（使用 JPEG 压缩以提速）"""
    try:
        with Image.open(image_path) as img:
            # 转为 RGB，防止 RGBA 报错
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            buf = io.BytesIO()
            # 使用 JPEG 质量 85，兼顾速度与画质
            img.save(buf, format='JPEG', quality=85)
            return base64.b64encode(buf.getvalue()).decode('utf-8')
    except Exception as e:
        print(f"图片编码失败 {image_path}: {e}")
        return None

def pdf_to_images(pdf_path, dpi=150):
    dpi = max(50, min(500, dpi))
    image_paths = []
    temp_dir = tempfile.mkdtemp(prefix="pdf_to_imgs_")
    _temp_dirs.append(temp_dir)
    try:
        with fitz.open(pdf_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                zoom = dpi / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
                pix.save(img_path)
                image_paths.append(img_path)
    except Exception as e:
        print(f"PDF转换错误: {e}")
    return image_paths

def update_preview(files, dpi=150):
    dpi = max(50, min(500, dpi))
    if not files:
        return []
    preview_images = []
    for file in files:
        file_path = file if isinstance(file, str) else file.name
        if file_path.lower().endswith('.pdf'):
            preview_images.extend(pdf_to_images(file_path, dpi))
        else:
            preview_images.append(file_path)
    return preview_images

def markdown_to_docx(md_text, output_path):
    """保留原版表格解析逻辑，将 Markdown 转为 Word"""
    doc = Document()
    lines = md_text.split('\n')
    in_table = False
    table_rows = []
    for line in lines:
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip('|').split('|')]
            if re.match(r'^[\s\-:]+$', ' '.join(cells)):
                continue
            table_rows.append(cells)
        else:
            if in_table and table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                for i, row_cells in enumerate(table_rows):
                    for j, cell_text in enumerate(row_cells):
                        if j < num_cols:
                            table.cell(i, j).text = cell_text
                doc.add_paragraph()
                in_table = False
                table_rows = []
            stripped = line.lstrip('#').strip()
            heading_level = len(line) - len(line.lstrip('#'))
            if 1 <= heading_level <= 4 and stripped:
                doc.add_heading(stripped, level=heading_level)
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)
            elif line.strip():
                doc.add_paragraph(line)
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        for i, row_cells in enumerate(table_rows):
            for j, cell_text in enumerate(row_cells):
                if j < num_cols:
                    table.cell(i, j).text = cell_text
    doc.save(output_path)

# ==================== 核心识别（调用 llama.cpp API） ====================

def batch_llama_ocr_recognition(files, dpi, progress=gr.Progress()):
    """改用 HTTP 请求，调用极速的 llama.cpp 服务"""
    if not files:
        yield "请至少上传一张图片或 PDF", "0.00s", ""
        return

    # 构建任务列表（支持 PDF 分页）
    tasks = []
    for file in files:
        file_path = file if isinstance(file, str) else file.name
        filename = os.path.basename(file_path)
        if file_path.lower().endswith('.pdf'):
            img_paths = pdf_to_images(file_path, dpi)
            for i, img_path in enumerate(img_paths):
                tasks.append({'img_path': img_path, 'filename': filename, 'is_pdf': True, 'page': i + 1})
        else:
            tasks.append({'img_path': file_path, 'filename': filename, 'is_pdf': False, 'page': None})

    total = len(tasks)
    if total == 0:
        yield "未找到可识别的内容", "0.00s", ""
        return

    all_results = []
    start_total = time.time()

    for i, task in enumerate(tasks):
        desc = f"识别: {task['filename']}" + (f" 第{task['page']}页" if task['is_pdf'] else "")
        progress(i / total, desc=desc)

        # 编码图片
        image_b64 = encode_image(task['img_path'])
        if not image_b64:
            header = f"文件: {task['filename']}" + (f" (第 {task['page']} 页)" if task['is_pdf'] else "")
            all_results.append(f"{header}\n\n错误: 图片编码失败\n\n{'─' * 50}\n")
            continue

        # 构建请求 (针对 MinerU 优化提示词，提升表格和公式输出)
        payload = {
            "model": MODEL_NAME,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "请识别图中的所有文字，将表格转换为标准Markdown表格格式，保留数学公式，输出纯净的Markdown文本。"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}}
                    ]
                }
            ],
            "max_tokens": MAX_TOKENS,
            "temperature": TEMPERATURE,
            "stream": False
        }

        img_start = time.time()
        try:
            resp = requests.post(LLAMA_API_URL, json=payload, timeout=120)
            resp.raise_for_status()
            data = resp.json()
            result = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            result = result.strip() if result else "未识别到内容"
        except requests.exceptions.ConnectionError:
            result = "连接失败，请确保 llama-server 已启动 (端口 50600)"
        except Exception as e:
            result = f"请求错误: {str(e)}"

        # 拼接结果
        header = f"文件: {task['filename']}" + (f" (第 {task['page']} 页)" if task['is_pdf'] else "")
        all_results.append(f"{header}\n\n{result}\n\n{'─' * 50}\n")

        img_cost = time.time() - img_start
        total_cost = time.time() - start_total
        combined = "\n".join(all_results)
        yield combined, f"单张: {img_cost:.2f}s | 总计: {total_cost:.2f}s", combined

    total_time = time.time() - start_total
    final_md = "\n".join(all_results)
    cleanup_temp_dirs()
    yield final_md, f"完成! 总耗时: {total_time:.2f}s", final_md


def export_word(markdown_text):
    if not markdown_text or "等待识别" in markdown_text or "请先上传" in markdown_text:
        return None
    temp_dir = tempfile.gettempdir()
    word_path = os.path.join(temp_dir, f"mineru_export_{int(time.time())}.docx")
    try:
        markdown_to_docx(markdown_text, word_path)
        return word_path
    except Exception as e:
        print(f"Word 导出失败: {e}")
        return None

# ==================== UI 构建（保留原版布局，微调文字） ====================

def create_ui():
    with gr.Blocks(title="MinerU-OCR 极速版 (llama.cpp)") as demo:

        gr.HTML(
            "<div style='text-align:center;padding:10px 0'>"
            "<h2 style='margin:0'>MinerU-OCR 极速版 (llama.cpp 后端)</h2>"
            "<p style='margin:4px 0 0;color:#666'>利用 llama.cpp 加速推理，支持表格/公式识别，导出 Word</p>"
            "</div>"
        )

        with gr.Row():
            # ---- 左栏 ----
            with gr.Column(scale=1):
                gr.HTML("<b>上传文件</b>")
                image_input = gr.File(
                    label="支持图片或 PDF（可多选）",
                    file_count="multiple",
                    file_types=["image", ".pdf"],
                    type="filepath"
                )
                preview_gallery = gr.Gallery(
                    label="预览", columns=[2], rows=[2],
                    object_fit="contain", height=400
                )

                gr.HTML("<b>推理设置</b>")
                # 修改：显示当前连接的服务器状态，不需要加载模型了
                model_status = gr.Textbox(
                    label="后端状态", 
                    value=f"已连接 {LLAMA_API_URL} (模型: {MODEL_NAME})", 
                    interactive=False, 
                    lines=1
                )

                dpi_slider = gr.Slider(
                    minimum=50, maximum=500, step=10,
                    value=150, label="PDF DPI (影响清晰度)"
                )

                with gr.Row():
                    submit_btn = gr.Button("开始识别", variant="primary")
                    # 移除加载/卸载模型按钮，因为现在直接调用 API
                    refresh_btn = gr.Button("测试连接", size="sm")

            # ---- 右栏 ----
            with gr.Column(scale=1):
                gr.HTML("<b>识别结果 (Markdown)</b>")
                output_md = gr.Markdown(
                    value="等待识别...\n\n确保 `llama-server` 已在端口 50600 运行。",
                    height=500
                )
                with gr.Row():
                    export_btn = gr.Button("导出 Word", variant="secondary", size="sm")
                    export_file = gr.File(label="下载", visible=False)

                gr.HTML("<b>耗时统计</b>")
                time_cost_box = gr.Textbox(
                    label="性能", lines=2, interactive=False
                )

        md_state = gr.State(value="")

        # ---- 事件绑定 ----
        refresh_btn.click(
            fn=lambda: f"连接测试中... (请求 {LLAMA_API_URL})",
            outputs=[model_status]
        ).then(
            fn=lambda: "连接成功 ✅" if requests.get(LLAMA_API_URL.replace("/v1/chat/completions", "/health"), timeout=2).ok else "连接失败 ❌ 请检查服务",
            outputs=[model_status]
        )

        image_input.change(
            fn=update_preview,
            inputs=[image_input, dpi_slider],
            outputs=[preview_gallery]
        )
        dpi_slider.change(
            fn=update_preview,
            inputs=[image_input, dpi_slider],
            outputs=[preview_gallery]
        )

        submit_btn.click(
            fn=batch_llama_ocr_recognition,
            inputs=[image_input, dpi_slider],
            outputs=[output_md, time_cost_box, md_state]
        )

        export_btn.click(
            fn=export_word,
            inputs=[md_state],
            outputs=[export_file]
        ).then(lambda: gr.update(visible=True), None, [export_file])

    return demo


if __name__ == "__main__":
    demo = create_ui()
    print("=" * 60)
    print("启动 MinerU-OCR 极速版 (llama.cpp 后端)")
    print(f"后端地址: {LLAMA_API_URL}")
    print("请确保已启动 llama-server (端口 50600)")
    print("=" * 60)
    demo.launch(server_name="127.0.0.1", server_port=7889, share=False, inbrowser=True)